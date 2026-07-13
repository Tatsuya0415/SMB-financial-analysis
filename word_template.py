"""
Wordレポート生成モジュール（10ページ構成）
コンサルティングレポート品質・中小企業診断士基準
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import copy


# ─────────────────────────────────────────
# カラー定義
# ─────────────────────────────────────────
NAVY    = RGBColor(0x1B, 0x3A, 0x5C)
BLUE    = RGBColor(0x31, 0x82, 0xCE)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
GRAY    = RGBColor(0x4A, 0x55, 0x68)
LGRAY   = RGBColor(0xE2, 0xE8, 0xF0)
GREEN   = RGBColor(0x27, 0x67, 0x49)
AMBER   = RGBColor(0xB7, 0x79, 0x1F)
RED     = RGBColor(0xC5, 0x30, 0x30)
GREEN_BG= RGBColor(0xC6, 0xF6, 0xD5)
AMBER_BG= RGBColor(0xFE, 0xFC, 0xBF)
RED_BG  = RGBColor(0xFE, 0xD7, 0xD7)
LIGHT_BG= RGBColor(0xEB, 0xF8, 0xFF)

RAG_COLOR_MAP = {
    "GREEN": (GREEN, GREEN_BG, "◎ 良好"),
    "AMBER": (AMBER, AMBER_BG, "△ 要注意"),
    "RED":   (RED,   RED_BG,   "✕ 要改善"),
}


# ─────────────────────────────────────────
# ヘルパー関数
# ─────────────────────────────────────────
def _set_cell_bg(cell, rgb: RGBColor):
    """セル背景色を設定"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tc_pr.append(shd)


def _set_cell_border(cell, color="CBD5E0", size=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    for side in ['top', 'left', 'bottom', 'right']:
        bd = OxmlElement(f'w:{side}')
        bd.set(qn('w:val'), 'single')
        bd.set(qn('w:sz'), str(size))
        bd.set(qn('w:color'), color)
        tc_pr.append(bd)


def _para(doc, text="", bold=False, size=10, color=None,
           align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0,
           italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return p


def _heading(doc, text, level=1, size=14, color=NAVY, space_before=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def _section_bar(doc, text, color=NAVY):
    """セクションタイトルバー"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

    # 左ボーダーライン風（■ マーカー）
    run = p.add_run(f"▌ {text}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = color
    return p


def _divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 60)
    run.font.size = Pt(7)
    run.font.color.rgb = LGRAY
    return p


def _page_break(doc):
    doc.add_page_break()


def _fmt(val, fmt=","):
    if val is None:
        return "─"
    if fmt == "%":
        return f"{val:.1%}"
    if fmt == "%+":
        return f"{val:+.1%}"
    if fmt == "x":
        return f"{val:.2f}倍"
    if fmt == "d":
        return f"{val:.0f}日"
    return f"{val:,.0f}"


def _rag_label(rag: str) -> str:
    return RAG_COLOR_MAP.get(rag, (GRAY, WHITE, "─"))[2]


def _rag_fc(rag: str) -> RGBColor:
    return RAG_COLOR_MAP.get(rag, (GRAY, WHITE, "─"))[0]


def _rag_bg(rag: str) -> RGBColor:
    return RAG_COLOR_MAP.get(rag, (GRAY, WHITE, "─"))[1]


# ─────────────────────────────────────────
# P1 — 表紙
# ─────────────────────────────────────────
def build_cover(doc: Document, meta: dict):
    # 大きめスペース
    for _ in range(6):
        doc.add_paragraph()

    # 会社名
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta.get("company_name", "〇〇株式会社"))
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY

    doc.add_paragraph()

    # レポートタイトル
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("財務分析・経営診断レポート")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"― {meta.get('fiscal_year', '当期')} ―")
    run.font.size = Pt(14)
    run.font.color.rgb = GRAY

    for _ in range(4):
        doc.add_paragraph()

    _divider(doc)

    # メタ情報
    for label, val in [
        ("業種",       meta.get("industry", "─")),
        ("決算期",     meta.get("fiscal_year", "─")),
        ("診断実施日", meta.get("report_date", str(date.today()))),
        ("作成",       "中小企業財務診断システム（中小企業診断士基準）"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}：{val}")
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY

    _page_break(doc)


# ─────────────────────────────────────────
# P2 — エグゼクティブサマリー
# ─────────────────────────────────────────
def build_executive_summary(doc: Document, ratios: dict, issues: list,
                              recommendations: list, benchmark: dict,
                              overall_rag: str):
    _heading(doc, "エグゼクティブサマリー", size=16)
    _divider(doc)

    # 総合評価バナー
    rag_fc, rag_bg, rag_label = RAG_COLOR_MAP.get(
        overall_rag, (GRAY, WHITE, "評価中"))

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, rag_bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"総合評価：{rag_label}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = rag_fc

    doc.add_paragraph()

    # KPIハイライト（3列）
    _section_bar(doc, "主要KPIサマリー")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"

    kpi_items = [
        ("営業利益率",  ratios.get("operating_margin"), "%",  ratios.get("_rag_operating_margin",  "─")),
        ("自己資本比率", ratios.get("equity_ratio"),    "%",  ratios.get("_rag_equity_ratio",       "─")),
        ("ROA",         ratios.get("roa"),              "%",  ratios.get("_rag_roa",                "─")),
        ("流動比率",    ratios.get("current_ratio"),    "x",  ratios.get("_rag_current_ratio",      "─")),
        ("売上高成長率",ratios.get("revenue_growth"),   "%+", ratios.get("_rag_revenue_growth",     "─")),
        ("CCC",         ratios.get("ccc"),              "d",  ratios.get("_rag_ccc",                "─")),
    ]

    # 2行 × 3列 KPIカード
    tbl = doc.add_table(rows=2, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, (label, val, fmt, rag) in enumerate(kpi_items[:6]):
        row_idx = idx // 3
        col_idx = idx % 3
        cell = tbl.cell(row_idx, col_idx)
        fc, bg, rag_txt = RAG_COLOR_MAP.get(rag, (GRAY, WHITE, "─"))
        _set_cell_bg(cell, bg)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)

        run = p.add_run(label + "\n")
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(4)
        run2 = p2.add_run(_fmt(val, fmt))
        run2.bold = True
        run2.font.size = Pt(16)
        run2.font.color.rgb = fc

        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(6)
        run3 = p3.add_run(rag_txt)
        run3.font.size = Pt(8)
        run3.font.color.rgb = fc

    doc.add_paragraph()

    # Top 3 課題
    _section_bar(doc, "特定された主要課題（Top 3）")
    top3 = (issues or [])[:3]
    for i, issue in enumerate(top3, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"{i}. {issue.get('title', '─')}")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY

        if issue.get("description"):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(1.2)
            p2.paragraph_format.space_after = Pt(4)
            run2 = p2.add_run(issue["description"])
            run2.font.size = Pt(9)
            run2.font.color.rgb = GRAY

    doc.add_paragraph()

    # 戦略提言サマリー
    _section_bar(doc, "主要提言（即時〜短期）")
    top_recs = (recommendations or [])[:3]
    for i, rec in enumerate(top_recs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"▶ {rec.get('action', '─')}")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BLUE

        if rec.get("expected_effect"):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(1.2)
            p2.paragraph_format.space_after = Pt(2)
            run2 = p2.add_run(f"期待効果: {rec['expected_effect']}")
            run2.font.size = Pt(9)
            run2.font.color.rgb = GRAY
            run2.italic = True

    _page_break(doc)


# ─────────────────────────────────────────
# P3-4 — 財務ハイライト
# ─────────────────────────────────────────
def build_financial_highlights(doc: Document, data: dict, ratios: dict,
                                 benchmark: dict):
    _heading(doc, "財務ハイライト", size=16)
    _divider(doc)

    pl = data.get("pl", {})
    bs = data.get("bs", {})

    # 主要財務数値テーブル
    _section_bar(doc, "損益・財務ポジション")

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["科目", "当期（千円）", "前期（千円）", "増減"]
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        _set_cell_bg(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    fin_rows = [
        ("売上高",     pl.get("revenue"),          pl.get("revenue_prev")),
        ("売上総利益", pl.get("gross_profit"),      pl.get("gross_profit_prev")),
        ("営業利益",   pl.get("operating_profit"),  pl.get("operating_profit_prev")),
        ("経常利益",   pl.get("ordinary_profit"),   pl.get("ordinary_profit_prev")),
        ("当期純利益", pl.get("net_profit"),         pl.get("net_profit_prev")),
        ("総資産",     bs.get("total_assets"),       bs.get("total_assets_prev")),
        ("純資産",     bs.get("equity"),             bs.get("equity_prev")),
    ]

    for i, (label, cur, prev) in enumerate(fin_rows):
        row = tbl.add_row()
        bg = LIGHT_BG if i % 2 == 0 else WHITE

        cell_label = row.cells[0]
        _set_cell_bg(cell_label, bg)
        p = cell_label.paragraphs[0]
        run = p.add_run(label)
        run.bold = label in ["売上高", "営業利益"]
        run.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)

        for j, val in enumerate([cur, prev], 1):
            c = row.cells[j]
            _set_cell_bg(c, bg)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(_fmt(val))
            run.font.size = Pt(10)

        diff_cell = row.cells[3]
        _set_cell_bg(diff_cell, bg)
        p = diff_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if cur is not None and prev is not None:
            diff = cur - prev
            run = p.add_run(f"{diff:+,.0f}")
            run.font.size = Pt(10)
            run.font.color.rgb = GREEN if diff >= 0 else RED

    doc.add_paragraph()

    # 主要指標テーブル（業界比較付き）
    _section_bar(doc, "主要経営指標（業界比較）")

    tbl2 = doc.add_table(rows=1, cols=5)
    tbl2.style = "Table Grid"

    for i, h in enumerate(["指標", "当社値", "業界平均", "過不足", "評価"]):
        cell = tbl2.cell(0, i)
        _set_cell_bg(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    kpi_rows = [
        ("営業利益率",   "operating_margin",  "%"),
        ("経常利益率",   "ordinary_margin",   "%"),
        ("ROA",         "roa",               "%"),
        ("ROE",         "roe",               "%"),
        ("自己資本比率", "equity_ratio",       "%"),
        ("流動比率",     "current_ratio",     "x"),
        ("CCC（日数）", "ccc",               "d"),
        ("売上高成長率", "revenue_growth",    "%+"),
    ]

    for i, (label, key, fmt) in enumerate(kpi_rows):
        row = tbl2.add_row()
        bg = LIGHT_BG if i % 2 == 0 else WHITE
        val = ratios.get(key)
        bm  = benchmark.get(key)
        rag = ratios.get(f"_rag_{key}", "─")

        _set_cell_bg(row.cells[0], bg)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(label)
        run.font.size = Pt(9)

        for j, v in enumerate([val, bm], 1):
            c = row.cells[j]
            _set_cell_bg(c, bg)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(_fmt(v, fmt))
            run.font.size = Pt(10)

        # 過不足
        diff_cell = row.cells[3]
        _set_cell_bg(diff_cell, bg)
        p = diff_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if val is not None and bm is not None:
            diff = val - bm
            run = p.add_run(f"{diff:+.1%}" if fmt in ["%", "%+"] else f"{diff:+.2f}")
            run.font.size = Pt(9)
            run.font.color.rgb = GREEN if diff >= 0 else RED

        # RAG
        rag_cell = row.cells[4]
        fc, bg_rag, rag_txt = RAG_COLOR_MAP.get(rag, (GRAY, WHITE, rag))
        _set_cell_bg(rag_cell, bg_rag)
        p = rag_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(rag_txt)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = fc

    _page_break(doc)


# ─────────────────────────────────────────
# P5 — 収益性分析
# ─────────────────────────────────────────
def build_profitability(doc: Document, ratios: dict, benchmark: dict,
                          pl: dict):
    _heading(doc, "収益性分析", size=16)
    _divider(doc)

    _section_bar(doc, "収益構造の概観")

    rev = pl.get("revenue") or 1
    gm  = pl.get("gross_profit", 0) or 0
    sga = pl.get("sga", 0) or 0
    op  = pl.get("operating_profit", 0) or 0
    oc  = pl.get("ordinary_profit", 0) or 0
    net = pl.get("net_profit", 0) or 0

    items = [
        ("売上高",     rev,      rev/rev,  NAVY),
        ("売上総利益", gm,       gm/rev,   BLUE),
        ("営業利益",   op,       op/rev,   GREEN),
        ("経常利益",   oc,       oc/rev,   GREEN),
        ("当期純利益", net,      net/rev,  GREEN if net >= 0 else RED),
    ]

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"

    for i, h in enumerate(["科目", "金額（千円）", "対売上比率", "業界平均比率"]):
        cell = tbl.cell(0, i)
        _set_cell_bg(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    key_map = [None, "gross_margin", "operating_margin", "ordinary_margin", None]
    for i, ((label, val, ratio, color), key) in enumerate(zip(items, key_map)):
        row = tbl.add_row()
        bg = LIGHT_BG if i % 2 == 0 else WHITE

        _set_cell_bg(row.cells[0], bg)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(label)
        run.font.size = Pt(9)
        run.bold = True

        for j, v in enumerate([val, ratio], 1):
            c = row.cells[j]
            _set_cell_bg(c, bg)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(_fmt(v, "%" if j == 2 else ","))
            run.font.size = Pt(10)
            run.font.color.rgb = color

        # 業界平均
        bm_cell = row.cells[3]
        _set_cell_bg(bm_cell, bg)
        p = bm_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        bm_val = benchmark.get(key) if key else None
        run = p.add_run(_fmt(bm_val, "%") if bm_val is not None else "─")
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY

    doc.add_paragraph()

    # 収益性の診断コメント
    _section_bar(doc, "診断コメント")
    op_margin = ratios.get("operating_margin", 0) or 0
    bm_op = benchmark.get("operating_margin", 0.03)
    gm_rate = (pl.get("gross_profit") or 0) / (pl.get("revenue") or 1)
    bm_gm = benchmark.get("gross_margin", 0.25)

    comments = []
    if gm_rate < bm_gm - 0.05:
        comments.append("【粗利率】業界平均を大きく下回っており、"
                        "価格設定の見直しまたは仕入原価の削減が急務です。")
    elif gm_rate > bm_gm + 0.05:
        comments.append("【粗利率】業界平均を大きく上回る高収益体質を確認しました。"
                        "競争優位の源泉を特定し、持続的拡大を検討してください。")
    else:
        comments.append("【粗利率】業界平均水準にあります。"
                        "粗利率改善施策により収益性向上の余地があります。")

    if op_margin < 0:
        comments.append("【営業利益率】営業赤字となっており、"
                        "コスト構造の抜本的見直しが最優先事項です。")
    elif op_margin < bm_op:
        comments.append("【営業利益率】業界平均を下回っています。"
                        "販管費の精査と生産性改善を進めてください。")
    else:
        comments.append("【営業利益率】業界平均を上回る収益力を維持しています。")

    roe = ratios.get("roe", 0) or 0
    if roe < 0.05:
        comments.append("【ROE】株主資本利益率が低水準。"
                        "利益成長と財務レバレッジ管理の両面からの改善が必要です。")

    for comment in comments:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.3)
        run = p.add_run(f"• {comment}")
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY

    _page_break(doc)


# ─────────────────────────────────────────
# P6 — 安全性・効率性分析
# ─────────────────────────────────────────
def build_safety_efficiency(doc: Document, ratios: dict, benchmark: dict):
    _heading(doc, "安全性・効率性分析", size=16)
    _divider(doc)

    # 安全性
    _section_bar(doc, "財務安全性")

    safety_items = [
        ("流動比率（≥120%が目安）",  "current_ratio",    "x"),
        ("当座比率（≥100%が目安）",  "quick_ratio",      "x"),
        ("自己資本比率（≥40%が目安）","equity_ratio",    "%"),
        ("負債比率（≤100%が目安）",  "debt_ratio",       "x"),
        ("インタレストカバレッジ",    "interest_coverage","x"),
        ("純有利子負債/EBITDA",      "net_debt_ebitda",  "x"),
    ]

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"

    for i, h in enumerate(["指標", "当社", "業界平均", "評価"]):
        cell = tbl.cell(0, i)
        _set_cell_bg(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    for i, (label, key, fmt) in enumerate(safety_items):
        row = tbl.add_row()
        bg = LIGHT_BG if i % 2 == 0 else WHITE
        val = ratios.get(key)
        bm = benchmark.get(key)
        rag = ratios.get(f"_rag_{key}", "─")

        _set_cell_bg(row.cells[0], bg)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(label)
        run.font.size = Pt(9)

        for j, v in enumerate([val, bm], 1):
            c = row.cells[j]
            _set_cell_bg(c, bg)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(_fmt(v, fmt))
            run.font.size = Pt(10)

        fc, bg_rag, rag_txt = RAG_COLOR_MAP.get(rag, (GRAY, WHITE, "─"))
        _set_cell_bg(row.cells[3], bg_rag)
        p = row.cells[3].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(rag_txt)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = fc

    doc.add_paragraph()

    # 効率性
    _section_bar(doc, "事業効率性（運転資本管理）")

    eff_items = [
        ("売掛金回転日数",     "receivable_days",  "d"),
        ("棚卸資産回転日数",   "inventory_days",   "d"),
        ("買掛金回転日数",     "payable_days",     "d"),
        ("CCC（現金転換サイクル）", "ccc",          "d"),
        ("総資産回転率",       "asset_turnover",   "x"),
    ]

    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Table Grid"

    for i, h in enumerate(["指標", "当社", "業界平均", "評価"]):
        cell = tbl2.cell(0, i)
        _set_cell_bg(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    for i, (label, key, fmt) in enumerate(eff_items):
        row = tbl2.add_row()
        bg = LIGHT_BG if i % 2 == 0 else WHITE
        val = ratios.get(key)
        bm = benchmark.get(key)
        rag = ratios.get(f"_rag_{key}", "─")

        _set_cell_bg(row.cells[0], bg)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(label)
        run.font.size = Pt(9)

        for j, v in enumerate([val, bm], 1):
            c = row.cells[j]
            _set_cell_bg(c, bg)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(_fmt(v, fmt))
            run.font.size = Pt(10)

        fc, bg_rag, rag_txt = RAG_COLOR_MAP.get(rag, (GRAY, WHITE, "─"))
        _set_cell_bg(row.cells[3], bg_rag)
        p = row.cells[3].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(rag_txt)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = fc

    # CCC コメント
    ccc = ratios.get("ccc")
    if ccc is not None:
        doc.add_paragraph()
        p = _para(doc, space_before=4, space_after=4)
        run = p.add_run(f"📌 CCC {_fmt(ccc, 'd')} — ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY
        if ccc > 60:
            run2 = p.add_run("運転資金の固定期間が長く、資金繰りへの影響が懸念されます。"
                              "在庫圧縮・回収サイト短縮を優先してください。")
        elif ccc < 20:
            run2 = p.add_run("現金転換サイクルが短く、資金効率は非常に良好な状態です。")
        else:
            run2 = p.add_run("業界水準範囲内ですが、継続的な最適化が収益力向上につながります。")
        run2.font.size = Pt(9)
        run2.font.color.rgb = GRAY

    _page_break(doc)


# ─────────────────────────────────────────
# P7 — 成長性・CF分析
# ─────────────────────────────────────────
def build_growth_cf(doc: Document, ratios: dict, data: dict):
    _heading(doc, "成長性・キャッシュフロー分析", size=16)
    _divider(doc)

    _section_bar(doc, "成長トレンド")

    growth_items = [
        ("売上高成長率",     "revenue_growth",   ratios.get("revenue_growth")),
        ("営業利益成長率",   "op_profit_growth", ratios.get("op_profit_growth")),
        ("自己資本成長率",   "equity_growth",    ratios.get("equity_growth")),
    ]

    for label, key, val in growth_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)

        run = p.add_run(f"{label}：")
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = NAVY

        if val is not None:
            color = GREEN if val >= 0 else RED
            run2 = p.add_run(f"{val:+.1%}")
            run2.font.size = Pt(14)
            run2.bold = True
            run2.font.color.rgb = color
        else:
            run2 = p.add_run("（前期データなし）")
            run2.font.size = Pt(10)
            run2.font.color.rgb = GRAY

    doc.add_paragraph()

    # CF分析
    cf = data.get("cf", {})
    if cf:
        _section_bar(doc, "キャッシュフロー分析")

        cf_items = [
            ("営業CF",  cf.get("operating_cf"),  "本業からの現金創出力"),
            ("投資CF",  cf.get("investing_cf"),  "設備投資・資産売却の動向"),
            ("財務CF",  cf.get("financing_cf"),  "借入・返済・配当の資金移動"),
        ]

        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"

        for i, h in enumerate(["区分", "金額（千円）", "意味"]):
            cell = tbl.cell(0, i)
            _set_cell_bg(cell, NAVY)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = WHITE

        for i, (label, val, desc) in enumerate(cf_items):
            row = tbl.add_row()
            bg = LIGHT_BG if i % 2 == 0 else WHITE

            _set_cell_bg(row.cells[0], bg)
            p = row.cells[0].paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(label)
            run.bold = True
            run.font.size = Pt(10)

            _set_cell_bg(row.cells[1], bg)
            p = row.cells[1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            if val is not None:
                run = p.add_run(f"{val:+,.0f}")
                run.font.size = Pt(11)
                run.font.color.rgb = GREEN if val >= 0 else RED
                run.bold = True

            _set_cell_bg(row.cells[2], bg)
            p = row.cells[2].paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(desc)
            run.font.size = Pt(8)
            run.font.color.rgb = GRAY
            run.italic = True

        # CF パターン診断
        doc.add_paragraph()
        op_cf = cf.get("operating_cf", 0) or 0
        inv_cf = cf.get("investing_cf", 0) or 0
        fin_cf = cf.get("financing_cf", 0) or 0

        pattern = ""
        if op_cf > 0 and inv_cf < 0 and fin_cf < 0:
            pattern = "【安定成長型】本業で稼ぎ、投資をしながら負債返済も進んでいます。健全な財務運営です。"
        elif op_cf > 0 and inv_cf < 0 and fin_cf > 0:
            pattern = "【積極投資型】本業収益＋外部調達で積極的な設備投資を行っています。成長投資の実効性を確認してください。"
        elif op_cf < 0:
            pattern = "【要注意】営業キャッシュフローがマイナスです。本業の収益力を至急改善する必要があります。"
        else:
            pattern = "CFパターンを評価中です。"

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.left_indent = Cm(0.3)
        run = p.add_run(f"💡 {pattern}")
        run.font.size = Pt(9)
        run.font.color.rgb = NAVY
        run.italic = True

    _page_break(doc)


# ─────────────────────────────────────────
# P8 — 課題の特定
# ─────────────────────────────────────────
def build_issues(doc: Document, issues: list):
    _heading(doc, "課題の特定と根本原因分析", size=16)
    _divider(doc)

    _para(doc,
          "MECEの原則（Mutually Exclusive, Collectively Exhaustive）に基づき、"
          "財務データから4つの視点で課題を特定します。",
          size=9, color=GRAY, space_before=0, space_after=8)

    for i, issue in enumerate(issues, 1):
        # 課題タイトル
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"課題 {i}：{issue.get('title', '─')}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY

        # 課題テーブル
        tbl = doc.add_table(rows=3, cols=2)
        tbl.style = "Table Grid"

        labels = ["現状（What）", "原因仮説（Why）", "財務的影響（Impact）"]
        keys   = ["current_state", "root_cause", "impact"]

        for j, (label, key) in enumerate(zip(labels, keys)):
            cell_l = tbl.cell(j, 0)
            cell_v = tbl.cell(j, 1)

            _set_cell_bg(cell_l, NAVY)
            p = cell_l.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(label)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = WHITE

            _set_cell_bg(cell_v, LIGHT_BG if j % 2 == 0 else WHITE)
            p = cell_v.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(issue.get(key, "─"))
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY

        doc.add_paragraph()

    _page_break(doc)


# ─────────────────────────────────────────
# P9-10 — 戦略提言ロードマップ
# ─────────────────────────────────────────
def build_recommendations(doc: Document, recommendations: list):
    _heading(doc, "戦略提言ロードマップ", size=16)
    _divider(doc)

    _para(doc,
          "優先度マトリクス（緊急度 × 重要度）に基づき、施策を3つのタイムラインに分類します。",
          size=9, color=GRAY, space_before=0, space_after=8)

    # タイムライン区分
    timeline_order = ["即時対応（0〜3ヶ月）", "短期施策（3〜12ヶ月）", "中長期施策（1〜3年）"]
    timeline_colors = [RED, AMBER, GREEN]

    by_timeline: dict = {}
    for rec in recommendations:
        tl = rec.get("timeline", "短期施策（3〜12ヶ月）")
        by_timeline.setdefault(tl, []).append(rec)

    for tl, color in zip(timeline_order, timeline_colors):
        recs = by_timeline.get(tl, [])
        if not recs:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"▌ {tl}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = color

        for rec in recs:
            # アクション
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(f"▶ {rec.get('action', '─')}")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = NAVY

            # 詳細テーブル
            tbl = doc.add_table(rows=1, cols=4)
            tbl.style = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            for i, h in enumerate(["施策詳細", "目標KPI", "期待効果", "実行体制"]):
                cell = tbl.cell(0, i)
                _set_cell_bg(cell, NAVY)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(h)
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = WHITE

            row = tbl.add_row()
            for i, (key, bg) in enumerate(zip(
                    ["detail", "kpi", "expected_effect", "owner"],
                    [WHITE, LIGHT_BG, WHITE, LIGHT_BG])):
                cell = row.cells[i]
                _set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(rec.get(key, "─"))
                run.font.size = Pt(9)
                run.font.color.rgb = GRAY

            doc.add_paragraph()

    # 総括
    _section_bar(doc, "まとめと今後の経営方針")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(
        "本診断では財務データに基づき、収益性・安全性・効率性・成長性の4軸から経営状態を分析しました。"
        "特定された課題に対する施策を着実に実行し、財務KPIのモニタリング体制を整備することで、"
        "持続的な企業価値向上が期待できます。"
        "定期的な財務診断（四半期ごと）の実施を推奨します。"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run.italic = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("※ 本レポートは中小企業診断士基準に準拠した財務診断システムによる自動生成です。")
    run.font.size = Pt(7)
    run.font.color.rgb = GRAY


# ─────────────────────────────────────────
# ページ設定
# ─────────────────────────────────────────
def setup_page(doc: Document):
    from docx.oxml.ns import qn
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # フォント既定
    doc.styles['Normal'].font.name = '游ゴシック'
    doc.styles['Normal'].font.size = Pt(10)
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '游ゴシック')


# ─────────────────────────────────────────
# メイン生成関数
# ─────────────────────────────────────────
def generate_word(meta: dict, data: dict, ratios: dict, benchmark: dict,
                   issues: list, recommendations: list, overall_rag: str,
                   output_path: str):
    doc = Document()
    setup_page(doc)

    build_cover(doc, meta)
    build_executive_summary(doc, ratios, issues, recommendations,
                             benchmark, overall_rag)
    build_financial_highlights(doc, data, ratios, benchmark)
    build_profitability(doc, ratios, benchmark, data.get("pl", {}))
    build_safety_efficiency(doc, ratios, benchmark)
    build_growth_cf(doc, ratios, data)
    build_issues(doc, issues)
    build_recommendations(doc, recommendations)

    doc.save(output_path)
    print(f"✓ Word生成完了: {output_path}")
